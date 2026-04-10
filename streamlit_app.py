"""
Concordia Pipeline v4 - Streamlit Application

Spec-driven clinical trial data harmonization with:
- SpecRegistry replaces RAG (near-instant load)
- Real-time progress tracking
- Provenance-first data lineage
- Interactive QC review
- Full traceability

Run with: streamlit run streamlit_app.py

Secrets Configuration:
    Create .streamlit/secrets.toml with:
        ANTHROPIC_API_KEY = "sk-ant-xxx..."  # Optional, for LLM features
"""

import streamlit as st
import pandas as pd
import os
import time
from pathlib import Path
from datetime import datetime
from typing import Dict
import json
import logging
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def load_secrets_to_env():
    """Load Streamlit secrets into environment variables."""
    secrets_loaded = []

    secret_mappings = {
        "ANTHROPIC_API_KEY": ["ANTHROPIC_API_KEY"],
    }

    try:
        for secret_name, env_names in secret_mappings.items():
            if secret_name in st.secrets:
                value = st.secrets[secret_name]
                for env_name in env_names:
                    os.environ[env_name] = value
                secrets_loaded.append(secret_name)
                logger.info(f"Loaded secret: {secret_name}")
    except Exception as e:
        logger.warning(f"Could not load secrets: {e}")

    return secrets_loaded


# Load secrets before importing pipeline components
secrets_loaded = load_secrets_to_env()

# Force light theme before any rendering
from streamlit import config as _stconfig
_stconfig.set_option("theme.base", "light")
_stconfig.set_option("theme.primaryColor", "#3B63A8")
_stconfig.set_option("theme.backgroundColor", "#FFFFFF")
_stconfig.set_option("theme.secondaryBackgroundColor", "#F7F8FA")
_stconfig.set_option("theme.textColor", "#1A1F3C")

# Page configuration
st.set_page_config(
    page_title="Concordia Pipeline v4",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ---------------------------------------------------------------------------
# CRDSA Brand Constants
# ---------------------------------------------------------------------------
BRAND_BLUE = "#3B63A8"
BRAND_PURPLE = "#8B4DAA"
BRAND_TEAL = "#207870"
BRAND_NAVY = "#1A1F3C"
BRAND_TEXT = "#3A3A3C"
BRAND_LIGHT_BG = "#F7F8FA"
BRAND_BORDER = "#E2E2E2"
BRAND_RED = "#C0392B"
BRAND_AMBER = "#D4A017"


def inject_brand_css():
    """Inject CRDSA brand CSS overrides for a light-themed, professional UI."""
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Barlow:wght@400;500;600&family=Barlow+Semi+Condensed:wght@600;700&display=swap');

    /* ---- Force light theme via Streamlit CSS variables ---- */
    :root, [data-testid="stAppViewContainer"],
    .stApp, [data-testid="stHeader"],
    [data-testid="stToolbar"],
    [data-testid="stDecoration"] {
        --primary-color: #3B63A8;
        --background-color: #FFFFFF;
        --secondary-background-color: #F7F8FA;
        --text-color: #1A1F3C;
    }
    .stApp {
        background-color: #FFFFFF !important;
    }
    [data-testid="stAppViewContainer"] {
        background-color: #FFFFFF !important;
    }
    [data-testid="stHeader"] {
        background-color: #FFFFFF !important;
    }
    [data-testid="stBottomBlockContainer"],
    [data-testid="stMainBlockContainer"] {
        background-color: #FFFFFF !important;
    }

    /* ---- Global font override ---- */
    html, body, [class*="css"] {
        font-family: 'Barlow', sans-serif;
        color: #3A3A3C;
        background-color: #FFFFFF;
    }

    /* ---- Headings ---- */
    h1, h2, h3 {
        font-family: 'Barlow Semi Condensed', sans-serif !important;
        color: #1A1F3C !important;
        font-weight: 600;
    }

    /* ---- Sidebar — dark navy ---- */
    section[data-testid="stSidebar"] {
        background-color: #1A1F3C;
    }
    section[data-testid="stSidebar"],
    section[data-testid="stSidebar"] * {
        color: #F0F0F1 !important;
    }
    section[data-testid="stSidebar"] hr {
        border-color: rgba(240,240,241,0.15);
    }
    section[data-testid="stSidebar"] [data-baseweb="select"] * {
        color: #F0F0F1 !important;
    }
    /* Sidebar select/dropdown inputs — dark bg so light text is readable */
    section[data-testid="stSidebar"] [data-baseweb="select"] > div {
        background-color: #2A3050 !important;
        border-color: rgba(240,240,241,0.2) !important;
    }
    section[data-testid="stSidebar"] [data-baseweb="select"] > div:hover {
        border-color: rgba(240,240,241,0.4) !important;
    }
    /* Dropdown menu (popover) — also dark */
    section[data-testid="stSidebar"] [data-baseweb="popover"] ul {
        background-color: #2A3050 !important;
    }
    section[data-testid="stSidebar"] [data-baseweb="popover"] li {
        color: #F0F0F1 !important;
    }
    section[data-testid="stSidebar"] [data-baseweb="popover"] li:hover {
        background-color: #3B63A8 !important;
    }

    /* ---- Primary button — brand blue ---- */
    .stButton > button[kind="primary"],
    .stButton > button[data-testid="stBaseButton-primary"] {
        background-color: #3B63A8 !important;
        border: none !important;
        color: white !important;
        font-family: 'Barlow Semi Condensed', sans-serif;
        font-weight: 600;
        letter-spacing: 0.02em;
        border-radius: 6px;
    }
    .stButton > button[kind="primary"]:hover,
    .stButton > button[data-testid="stBaseButton-primary"]:hover {
        background-color: #2E4F8A !important;
    }

    /* ---- Secondary / download buttons — ghost style ---- */
    .stDownloadButton > button,
    .stButton > button:not([kind="primary"]):not([data-testid="stBaseButton-primary"]) {
        background-color: transparent !important;
        border: 1.5px solid #3B63A8 !important;
        color: #3B63A8 !important;
        font-family: 'Barlow Semi Condensed', sans-serif;
        font-weight: 600;
        border-radius: 6px;
    }
    .stDownloadButton > button:hover,
    .stButton > button:not([kind="primary"]):hover {
        background-color: #EDF2FA !important;
    }

    /* ---- Tabs — brand blue underline ---- */
    .stTabs [data-baseweb="tab-list"] {
        gap: 0;
        border-bottom: 2px solid #E2E2E2;
    }
    .stTabs [data-baseweb="tab"] {
        font-family: 'Barlow Semi Condensed', sans-serif;
        font-weight: 500;
        color: #3A3A3C;
        padding: 10px 20px;
    }
    .stTabs [aria-selected="true"] {
        color: #3B63A8 !important;
        border-bottom: 3px solid #3B63A8;
        font-weight: 600;
    }

    /* ---- Metric cards ---- */
    [data-testid="stMetric"] {
        background-color: #F7F8FA;
        border: 1px solid #E2E2E2;
        border-radius: 8px;
        padding: 12px 16px;
    }

    /* ---- Progress bar — brand gradient ---- */
    .stProgress > div > div > div {
        background: linear-gradient(135deg, #3B63A8 0%, #8B4DAA 100%) !important;
    }

    /* ---- Expander headers ---- */
    .streamlit-expanderHeader,
    [data-testid="stExpander"] summary {
        font-family: 'Barlow Semi Condensed', sans-serif;
        font-weight: 600;
        color: #1A1F3C;
    }

    /* ---- Alert accent colors ---- */
    .stAlert [data-testid="stNotificationContentSuccess"] { border-left-color: #207870 !important; }
    .stAlert [data-testid="stNotificationContentError"]   { border-left-color: #C0392B !important; }
    .stAlert [data-testid="stNotificationContentWarning"] { border-left-color: #D4A017 !important; }
    .stAlert [data-testid="stNotificationContentInfo"]    { border-left-color: #3B63A8 !important; }

    /* ---- Dividers ---- */
    hr { border-color: #E2E2E2 !important; }

    /* ---- File uploader area ---- */
    [data-testid="stFileUploader"] {
        border-radius: 8px;
    }

    /* ---- Dataframe container ---- */
    [data-testid="stDataFrame"] {
        border: 1px solid #E2E2E2;
        border-radius: 8px;
    }
    </style>
    """, unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Brand UI Components
# ---------------------------------------------------------------------------

def status_badge(label: str, level: str = "info") -> str:
    """Return an HTML badge matching CRDSA brand colors."""
    colors = {
        "success": (BRAND_TEAL, "#E8F5F3"),
        "error":   (BRAND_RED, "#FDEDEB"),
        "warning": (BRAND_AMBER, "#FFF8E1"),
        "info":    (BRAND_BLUE, "#EDF2FA"),
        "purple":  (BRAND_PURPLE, "#F5EDF8"),
        "neutral": ("#6B7280", "#F3F4F6"),
    }
    fg, bg = colors.get(level, colors["info"])
    return (
        f'<span style="background:{bg}; color:{fg}; padding:4px 12px; '
        f'border-radius:4px; font-family:Barlow Semi Condensed,sans-serif; '
        f'font-weight:600; font-size:0.85rem; display:inline-block;">{label}</span>'
    )


def metric_card(label: str, value: str, accent: str = BRAND_BLUE,
                value_color: str = "") -> str:
    """Render a metric as a branded card with colored top-border accent."""
    val_color = value_color or BRAND_NAVY
    return f"""
    <div style="
        background: white;
        border: 1px solid {BRAND_BORDER};
        border-top: 3px solid {accent};
        border-radius: 8px;
        padding: 16px 20px;
        text-align: center;
        min-height: 90px;
        display: flex;
        flex-direction: column;
        justify-content: center;
    ">
        <p style="margin:0; font-size:0.75rem; color:#6B7280; font-family:Barlow,sans-serif;
                   text-transform:uppercase; letter-spacing:0.06em;">{label}</p>
        <p style="margin:6px 0 0; font-size:1.35rem; font-weight:700;
                   font-family:Barlow Semi Condensed,sans-serif; color:{val_color};">{value}</p>
    </div>
    """


def stoplight_panel(level: str, reason: str = "",
                    core_present: list = None, core_missing: list = None) -> str:
    """Render LLM review stoplight as a professional left-border panel."""
    color_map = {"GREEN": BRAND_TEAL, "YELLOW": BRAND_AMBER, "RED": BRAND_RED}
    bg_map = {"GREEN": "#E8F5F3", "YELLOW": "#FFF8E1", "RED": "#FDEDEB"}
    level_upper = level.upper()
    color = color_map.get(level_upper, BRAND_BLUE)
    bg = bg_map.get(level_upper, "#EDF2FA")

    html = f"""
    <div style="background:{bg}; border-left:4px solid {color};
                border-radius:0 8px 8px 0; padding:20px; margin:12px 0;">
        <div style="display:flex; align-items:center; gap:10px; margin-bottom:8px;">
            <span style="width:18px; height:18px; border-radius:50%;
                         background:{color}; display:inline-block;
                         box-shadow:0 0 8px {color}88;"></span>
            <span style="font-family:Barlow Semi Condensed,sans-serif; font-weight:700;
                         font-size:1.1rem; color:{color};">
                REVIEW: {level_upper}
            </span>
        </div>
    """
    if reason:
        html += f'<p style="margin:4px 0 0; color:{BRAND_TEXT}; font-family:Barlow,sans-serif;">{reason}</p>'

    if core_present:
        html += (
            f'<p style="margin:10px 0 2px; font-size:0.8rem; color:#6B7280; '
            f'text-transform:uppercase; letter-spacing:0.05em;">Core Variables Present</p>'
            f'<p style="margin:0; color:{BRAND_NAVY}; font-weight:500;">{", ".join(core_present)}</p>'
        )
    if core_missing:
        html += (
            f'<p style="margin:10px 0 2px; font-size:0.8rem; color:#6B7280; '
            f'text-transform:uppercase; letter-spacing:0.05em;">Core Variables Missing</p>'
            f'<p style="margin:0; color:{BRAND_RED}; font-weight:500;">{", ".join(core_missing)}</p>'
        )
    html += "</div>"
    return html


def pipeline_stepper(stages: list, current_index: int, failed: bool = False) -> str:
    """Render a horizontal pipeline stage stepper with state-driven colors.

    States:
    - Pending (not yet reached): gray with empty circle
    - Active (currently running): brand blue with spinner dot
    - Complete (successfully passed): teal with checkmark
    - Failed: red with X mark
    """
    html = '<div style="display:flex; gap:6px; margin:16px 0;">'
    for i, stage in enumerate(stages):
        if failed and i == current_index:
            # Failed at this stage
            color, bg, border = BRAND_RED, "#FDEDEB", BRAND_RED
            icon = "✗"
        elif i < current_index:
            # Completed
            color, bg, border = BRAND_TEAL, "#E8F5F3", BRAND_TEAL
            icon = "✓"
        elif i == current_index:
            # Active / in-progress
            color, bg, border = BRAND_BLUE, "#EDF2FA", BRAND_BLUE
            icon = "●"
        else:
            # Pending
            color, bg, border = "#9CA3AF", "#F9FAFB", "#E5E7EB"
            icon = "○"
        html += f'''
        <div style="flex:1; text-align:center; padding:8px 4px;
                    background:{bg}; border-radius:6px; border:1.5px solid {border};">
            <div style="font-size:0.85rem; margin-bottom:2px;">{icon}</div>
            <span style="font-family:Barlow Semi Condensed,sans-serif; font-weight:600;
                         font-size:0.7rem; color:{color}; letter-spacing:0.04em;
                         text-transform:uppercase;">{stage}</span>
        </div>'''
    html += '</div>'
    return html


# Import pipeline components
try:
    from orchestrator import PipelineOrchestrator, create_orchestrator, PipelineResult
    from config.settings import get_settings, reset_settings
    PIPELINE_AVAILABLE = True
except ImportError as e:
    PIPELINE_AVAILABLE = False
    IMPORT_ERROR = str(e)


def parse_data_dictionary(df: pd.DataFrame) -> Dict:
    """
    Parse a data dictionary DataFrame using deterministic rules.

    Output format:
    {
        "SEX": {"codes": {"1": "Male", "2": "Female"}},
        "RACE": {"codes": {"11": "White", "12": "Black or African American"}}
    }
    """
    dictionary = {}

    # Find the header row by looking for known column headers
    header_row_idx = None
    for idx, row in df.iterrows():
        for col_idx, cell in enumerate(row):
            if pd.notna(cell):
                cell_str = str(cell).strip().upper()
                if cell_str in ['VARIABLE NAME', 'VARIABLE', 'VAR NAME']:
                    header_row_idx = idx
                    break
        if header_row_idx is not None:
            break

    # If we found a header row, reindex the DataFrame
    if header_row_idx is not None:
        new_columns = df.iloc[header_row_idx].tolist()
        df = df.iloc[header_row_idx + 1:].copy()
        df.columns = [str(c).strip() if pd.notna(c) else f'col_{i}' for i, c in enumerate(new_columns)]

    # Identify key columns by mapping column names (case-insensitive)
    col_map = {str(c).upper().replace('\n', ' '): c for c in df.columns}

    var_col = None
    value_col = None
    format_col = None

    for candidate in ['VARIABLE NAME', 'VARIABLE', 'VAR', 'NAME', 'FIELD']:
        if candidate in col_map:
            var_col = col_map[candidate]
            break

    for candidate in ['VALID VALUES', 'VALUES', 'DECODE', 'VALID VALUE']:
        if candidate in col_map:
            value_col = col_map[candidate]
            break

    for candidate in ['FORMAT  (VALUE LIST)', 'FORMAT (VALUE LIST)', 'FORMAT', 'VALUE LIST', 'CODELIST']:
        if candidate in col_map:
            format_col = col_map[candidate]
            break

    if not var_col:
        logger.warning("Could not find variable name column in dictionary")
        return {}

    # Parse the dictionary - track current variable for continuation rows
    current_var = None

    for _, row in df.iterrows():
        var_name = row.get(var_col)

        if pd.notna(var_name) and str(var_name).strip():
            current_var = str(var_name).strip().upper()
            if current_var not in dictionary:
                dictionary[current_var] = {
                    "codes": {},
                    "format": str(row.get(format_col, '')) if format_col and pd.notna(row.get(format_col)) else ''
                }

        if current_var and value_col:
            value_str = row.get(value_col)
            if pd.notna(value_str):
                value_str = str(value_str).strip()
                if '=' in value_str:
                    parts = value_str.split('=', 1)
                    code = parts[0].strip()
                    label = parts[1].strip()
                    dictionary[current_var]["codes"][code] = label

    # Remove variables with no codes
    dictionary = {k: v for k, v in dictionary.items() if v.get("codes")}

    logger.info(f"Deterministic parser found {len(dictionary)} variables with codes")
    for var, data in dictionary.items():
        codes = data.get("codes", {})
        logger.info(f"  {var}: {len(codes)} codes - keys: {list(codes.keys())}")

    return dictionary


def parse_pdf_dictionary(pdf_file) -> Dict:
    """
    Parse a PDF data dictionary using pdfplumber table extraction.

    Handles two common formats:
    1. Tabular dictionaries (e.g., Amgen_265) with structured tables containing
       Variable Name, Variable Label, Type/Length, Decodes/Format, Origin, etc.
    2. Text-only PDFs (e.g., Merck_188) that contain no extractable tables.

    Returns the same dict format as parse_data_dictionary:
    {
        "SEX": {"codes": {"1": "Male", "2": "Female"}},
        "RACE": {"codes": {"11": "White", "12": "Black or African American"}}
    }
    """
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed — cannot parse PDF dictionaries")
        return {}

    dictionary = {}
    all_rows = []

    try:
        with pdfplumber.open(pdf_file) as pdf:
            logger.info(f"PDF has {len(pdf.pages)} pages")

            for page_num, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        all_rows.extend(table)

            if not all_rows:
                # Fallback: try extracting text and parsing with regex
                logger.info("No tables found in PDF — trying text extraction")
                full_text = ""
                for page in pdf.pages:
                    full_text += (page.extract_text() or "") + "\n"

                if not full_text.strip():
                    logger.warning("PDF contains no extractable text or tables")
                    return {}

                return _parse_pdf_text_dictionary(full_text)

    except Exception as e:
        logger.exception(f"Failed to read PDF: {e}")
        return {}

    if not all_rows:
        return {}

    # Identify header row and column positions
    header_idx, col_indices = _find_pdf_table_headers(all_rows)

    if header_idx is None or not col_indices.get("var_name"):
        logger.warning("Could not identify variable name column in PDF tables")
        # Try text-based fallback
        try:
            with pdfplumber.open(pdf_file) as pdf:
                full_text = ""
                for page in pdf.pages:
                    full_text += (page.extract_text() or "") + "\n"
                return _parse_pdf_text_dictionary(full_text)
        except Exception:
            return {}

    # Parse rows after header
    data_rows = all_rows[header_idx + 1:]
    var_col = col_indices["var_name"]
    decode_col = col_indices.get("decodes")
    label_col = col_indices.get("label")
    comment_col = col_indices.get("comments")

    current_var = None

    for row in data_rows:
        if not row or len(row) <= var_col:
            continue

        # Get variable name
        var_cell = row[var_col] if var_col < len(row) else None
        if var_cell and str(var_cell).strip():
            var_name = str(var_cell).strip().upper()
            # Skip domain headers (e.g., "AE.xpt", "DM.xpt", section titles)
            if var_name.endswith('.XPT') or var_name.startswith('DATASET'):
                continue
            # Skip if it looks like a column header repeated
            if var_name in ('VARIABLE NAME', 'VARIABLE', 'VAR NAME', 'NAME'):
                continue
            current_var = var_name

        if not current_var:
            continue

        if current_var not in dictionary:
            dictionary[current_var] = {"codes": {}}

        # Extract codes from Decodes/Format column
        if decode_col is not None and decode_col < len(row):
            decode_cell = row[decode_col]
            if decode_cell and str(decode_cell).strip():
                codes = _extract_codes_from_cell(str(decode_cell).strip())
                dictionary[current_var]["codes"].update(codes)

        # Also check Comments column for decode references
        if comment_col is not None and comment_col < len(row):
            comment_cell = row[comment_col]
            if comment_cell and str(comment_cell).strip():
                comment_str = str(comment_cell).strip()
                # Comments like "Decode of SEXCD" don't contain codes themselves
                # but label column might have the allowed values
                if label_col is not None and label_col < len(row):
                    label_cell = row[label_col]
                    if label_cell:
                        dictionary[current_var]["label"] = str(label_cell).strip()

    # Remove variables with no codes
    dictionary = {k: v for k, v in dictionary.items() if v.get("codes")}

    # Post-processing: filter out noise entries
    cleaned = {}
    for var, data in dictionary.items():
        # Skip repeated header rows that slipped through
        if re.search(r'\s', var) and not re.search(r'[_]', var):
            # Multi-word without underscore is likely a section header, not a variable
            continue

        codes = data.get("codes", {})

        # Remove SAS format noise from code values
        # Pattern: keys like "Follow", "$FORMAT.", "8.", "8.3", "(0", "1)"
        noise_keys = {k for k in codes if re.match(r'^[\d.()]+$', k)
                      or k.startswith('$')
                      or k.lower() == 'follow'
                      or re.match(r'^follow\s', k, re.IGNORECASE)}
        clean_codes = {k: v for k, v in codes.items() if k not in noise_keys}

        if clean_codes:
            data["codes"] = clean_codes
            cleaned[var] = data

    dictionary = cleaned

    logger.info(f"PDF parser found {len(dictionary)} variables with codes")
    for var, data in dictionary.items():
        codes = data.get("codes", {})
        logger.info(f"  {var}: {len(codes)} codes - keys: {list(codes.keys())[:10]}")

    return dictionary


def _find_pdf_table_headers(rows):
    """
    Find the header row in PDF table rows and return column indices.

    Returns (header_row_index, col_indices_dict) where col_indices_dict maps:
    - var_name: column index for Variable Name
    - label: column index for Variable Label
    - decodes: column index for Decodes / Format / Valid Values
    - comments: column index for Comments
    """
    header_keywords = {
        "var_name": ["VARIABLE NAME", "VARIABLE", "VAR NAME", "NAME", "FIELD"],
        "label": ["VARIABLE LABEL", "LABEL", "DESCRIPTION"],
        "decodes": [
            "DECODES / FORMAT", "DECODES/FORMAT", "DECODES",
            "FORMAT (VALUE LIST)", "FORMAT  (VALUE LIST)",
            "VALID VALUES", "VALUES", "CODELIST", "FORMAT",
            "DECODE", "VALUE LIST",
        ],
        "comments": ["COMMENTS", "COMMENT", "NOTES", "NOTE"],
    }

    for idx, row in enumerate(rows):
        if not row:
            continue
        # Normalize cells: uppercase, collapse whitespace/newlines
        cells_upper = []
        for c in row:
            if c:
                normalized = re.sub(r'\s+', ' ', str(c).strip().upper())
                cells_upper.append(normalized)
            else:
                cells_upper.append("")

        col_indices = {}
        for key, candidates in header_keywords.items():
            for col_idx, cell in enumerate(cells_upper):
                if cell in candidates:
                    col_indices[key] = col_idx
                    break

        if "var_name" in col_indices:
            logger.info(f"Found PDF header at row {idx}: {col_indices}")
            return idx, col_indices

    return None, {}


def _extract_codes_from_cell(text: str) -> Dict[str, str]:
    """
    Extract code-decode pairs from a Decodes/Format cell.

    Handles formats like:
    - "1=Male, 2=Female"
    - "Y, N"
    - "CHEILITIS, EYE, HAIR, NAIL, SKIN"
    - "1 = Male 2 = Female"
    - "M=Male F=Female"
    - "Follow $FREQ" (SAS format reference — skip)
    """
    codes = {}

    # Skip SAS format references like "$FREQ", "$SCOPE", "follow BODSYS."
    if text.startswith('$') or text.startswith('Follow $'):
        return codes

    # Skip SAS numeric formats like "8.", "8.3", "6.", "best12."
    if re.match(r'^(?:best)?\d+\.?\d*$', text, re.IGNORECASE):
        return codes

    # Skip "follow FORMAT." patterns (SAS format catalog references)
    if re.match(r'^(?:follow\s+)?\w+\.\s*$', text, re.IGNORECASE):
        return codes

    # Skip mixed SAS format + follow references like "8., follow\nBODSYS."
    cleaned = re.sub(r'\d+\.?\d*\s*,?\s*', '', text).strip()
    if re.match(r'^(?:follow\s+)?\w+\.?\s*$', cleaned, re.IGNORECASE) and '=' not in text:
        return codes

    # Try code=decode pairs first (e.g., "1=Male, 2=Female" or "M=Male\nF=Female")
    pairs = re.findall(r'([^,=\n]+?)\s*=\s*([^,=\n]+)', text)
    if pairs:
        for code, label in pairs:
            code = code.strip()
            label = label.strip()
            if code and label:
                codes[code] = label
        return codes

    # Try comma-separated list of values (e.g., "Y, N" or "CHEILITIS, EYE, HAIR")
    if ',' in text:
        values = [v.strip() for v in text.split(',') if v.strip()]
        # Only treat as value list if items are short (< 40 chars) — avoid capturing descriptions
        if values and all(len(v) < 40 for v in values):
            for v in values:
                codes[v] = v
            return codes

    # Try newline-separated values
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if len(lines) > 1 and all(len(l) < 40 for l in lines):
        for l in lines:
            # Check for "code = label" pattern in each line
            m = re.match(r'^(\S+)\s*=\s*(.+)', l)
            if m:
                codes[m.group(1).strip()] = m.group(2).strip()
            else:
                codes[l] = l

    return codes


def _parse_pdf_text_dictionary(text: str) -> Dict:
    """
    Fallback: parse a PDF dictionary from raw extracted text.

    Looks for patterns like:
    - "VARIABLE_NAME: 1=Male, 2=Female"
    - Lines with known variable names followed by code patterns
    """
    dictionary = {}

    # Known SDTM/clinical variable names to look for
    known_vars = [
        "SEX", "SEXCD", "RACE", "RACECD", "ETHNIC", "ETHNICCD",
        "COUNTRY", "ARMCD", "ARM", "AGEGP", "AGEGPCD", "AGEU",
        "AESER", "AEREL", "AEACN", "AEOUT", "AESEV", "AETOXGR",
        "AESCOPE", "AEFREQ",
    ]

    for var in known_vars:
        # Pattern: VAR followed by code list on same or next line
        pattern = rf'\b{var}\b[:\s]+((?:\d+\s*=\s*[^\n]+)(?:[,;\n]\s*\d+\s*=\s*[^\n]+)*)'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            codes_text = match.group(1)
            codes = _extract_codes_from_cell(codes_text)
            if codes:
                dictionary[var] = {"codes": codes}

    logger.info(f"PDF text parser found {len(dictionary)} variables with codes")
    return dictionary


def init_session_state():
    """Initialize session state variables."""
    if "pipeline_result" not in st.session_state:
        st.session_state.pipeline_result = None
    if "progress_log" not in st.session_state:
        st.session_state.progress_log = []
    if "current_stage" not in st.session_state:
        st.session_state.current_stage = None
    if "orchestrator" not in st.session_state:
        st.session_state.orchestrator = None


def progress_callback(stage: str, status: str, message: str, progress: float):
    """Legacy callback for pipeline progress updates (kept for compatibility)."""
    st.session_state.current_stage = stage
    st.session_state.progress_log.append({
        "timestamp": datetime.now().isoformat(),
        "stage": stage,
        "status": status,
        "message": message,
        "progress": progress
    })


def _render_knowledge_base_tree():
    """Render the knowledge base file tree in the sidebar.

    Shows each domain as an expander with categorized spec files:
    - Variable specs (DM_SEX.md, DM_AGE.md, ...)
    - Domain rules & plans
    - Value sets
    Clicking a file shows a preview snippet.
    """
    kb_path = Path(__file__).parent / "knowledge_base"
    if not kb_path.exists():
        st.caption("knowledge_base/ not found")
        return

    # Collect domains (subdirectories)
    domains = sorted([d.name for d in kb_path.iterdir() if d.is_dir()])

    # Also show top-level files (e.g., system_rules.md)
    top_files = sorted([f.name for f in kb_path.iterdir() if f.is_file() and f.suffix == '.md'])

    for domain in domains:
        domain_path = kb_path / domain
        md_files = sorted([f for f in domain_path.iterdir() if f.is_file() and f.suffix == '.md'])
        value_sets_path = domain_path / "value_sets"
        vs_files = sorted([f for f in value_sets_path.iterdir() if f.is_file() and f.suffix == '.md']) if value_sets_path.exists() else []

        # Categorize files
        var_specs = []
        domain_rules = []
        for f in md_files:
            name = f.name
            if 'domain_rules' in name:
                domain_rules.append(f)
            else:
                var_specs.append(f)

        total = len(var_specs) + len(domain_rules) + len(vs_files)

        with st.expander(f"**{domain}** — {total} specs", expanded=False):
            preview_key = f"kb_preview_{domain}"
            selected_path = st.session_state.get(preview_key, "")

            # Variable specs
            if var_specs:
                st.markdown(
                    f'<p style="font-size:0.75rem; color:{BRAND_TEAL}; margin:4px 0 2px; '
                    f'font-weight:600;">Variable Specs ({len(var_specs)})</p>',
                    unsafe_allow_html=True,
                )
                for f in var_specs:
                    var_name = f.stem.replace(f"{domain}_", "")
                    is_selected = str(f) == selected_path
                    btn_type = "primary" if is_selected else "secondary"
                    if st.button(f"📄 {var_name}", key=f"kb_{domain}_{f.stem}", use_container_width=True, type=btn_type):
                        st.session_state[preview_key] = str(f)
                        st.rerun()

            # Domain rules
            if domain_rules:
                st.markdown(
                    f'<p style="font-size:0.75rem; color:{BRAND_PURPLE}; margin:8px 0 2px; '
                    f'font-weight:600;">Domain Rules ({len(domain_rules)})</p>',
                    unsafe_allow_html=True,
                )
                for f in domain_rules:
                    label = f.stem.replace(f"{domain}_", "").replace("_", " ")
                    is_selected = str(f) == selected_path
                    btn_type = "primary" if is_selected else "secondary"
                    if st.button(f"📋 {label}", key=f"kb_{domain}_{f.stem}", use_container_width=True, type=btn_type):
                        st.session_state[preview_key] = str(f)
                        st.rerun()

            # Value sets
            if vs_files:
                st.markdown(
                    f'<p style="font-size:0.75rem; color:{BRAND_BLUE}; margin:8px 0 2px; '
                    f'font-weight:600;">Value Sets ({len(vs_files)})</p>',
                    unsafe_allow_html=True,
                )
                for f in vs_files:
                    label = f.stem.replace("_values", "").replace("_", " ").title()
                    is_selected = str(f) == selected_path
                    btn_type = "primary" if is_selected else "secondary"
                    if st.button(f"📊 {label}", key=f"kb_vs_{domain}_{f.stem}", use_container_width=True, type=btn_type):
                        st.session_state[preview_key] = str(f)
                        st.rerun()

            # Preview pane
            if preview_key in st.session_state and st.session_state[preview_key]:
                preview_path = Path(st.session_state[preview_key])
                if preview_path.exists():
                    st.divider()
                    content = preview_path.read_text(encoding="utf-8")
                    # Show first ~40 lines as a preview
                    lines = content.split("\n")
                    preview_text = "\n".join(lines[:40])
                    if len(lines) > 40:
                        preview_text += f"\n\n*... ({len(lines) - 40} more lines)*"
                    st.markdown(
                        f'<div style="font-size:0.75rem; color:#C0C8D8; background:#1A1F3C; '
                        f'padding:8px; border-radius:4px; max-height:300px; overflow-y:auto; '
                        f'white-space:pre-wrap; font-family:monospace;">{preview_text}</div>',
                        unsafe_allow_html=True,
                    )

    # Top-level files
    if top_files:
        with st.expander(f"**System** — {len(top_files)} files", expanded=False):
            sys_selected = st.session_state.get("kb_preview_system", "")
            for fname in top_files:
                label = fname.replace(".md", "").replace("_", " ").title()
                fpath = kb_path / fname
                is_selected = str(fpath) == sys_selected
                btn_type = "primary" if is_selected else "secondary"
                if st.button(f"⚙️ {label}", key=f"kb_sys_{fname}", use_container_width=True, type=btn_type):
                    st.session_state["kb_preview_system"] = str(fpath)
                    st.rerun()

            if "kb_preview_system" in st.session_state and st.session_state["kb_preview_system"]:
                sys_path = Path(st.session_state["kb_preview_system"])
                if sys_path.exists():
                    st.divider()
                    content = sys_path.read_text(encoding="utf-8")
                    lines = content.split("\n")
                    preview_text = "\n".join(lines[:40])
                    if len(lines) > 40:
                        preview_text += f"\n\n*... ({len(lines) - 40} more lines)*"
                    st.markdown(
                        f'<div style="font-size:0.75rem; color:#C0C8D8; background:#1A1F3C; '
                        f'padding:8px; border-radius:4px; max-height:300px; overflow-y:auto; '
                        f'white-space:pre-wrap; font-family:monospace;">{preview_text}</div>',
                        unsafe_allow_html=True,
                    )


def render_sidebar():
    """Render the sidebar with configuration options."""
    with st.sidebar:
        # Branded sidebar header
        st.markdown("""
        <div style="text-align:center; padding:4px 0 20px;">
            <p style="font-family:Barlow Semi Condensed,sans-serif; font-weight:700;
                      font-size:1.25rem; color:#F0F0F1; margin:0; letter-spacing:0.04em;">CONCORDIA</p>
            <p style="font-size:0.75rem; color:#8B9DC3; margin:2px 0 0;
                      font-family:Barlow,sans-serif;">Pipeline v4 — Spec-Driven</p>
        </div>
        """, unsafe_allow_html=True)

        st.subheader("LLM Settings")
        use_llm = st.checkbox(
            "Enable LLM",
            value=True,
            help="Use Claude for value resolution and review"
        )

        enable_review = st.checkbox(
            "Enable LLM Review",
            value=True,
            disabled=not use_llm,
            help="Run LLM-powered quality review stage"
        )

        st.subheader("Pipeline Options")
        domain = st.selectbox(
            "Domain",
            ["DM", "AE (preview)"],
            index=0,
            help="Target SDTM domain. AE requires domain specs in knowledge_base/AE/.",
            format_func=lambda x: x,
        )

        skip_qc = st.checkbox(
            "Skip QC Stage",
            value=False,
            help="Skip quality control checks"
        )

        st.divider()

        st.markdown(f"""
        <div style="font-family:Barlow,sans-serif; font-size:0.85rem; color:#C0C8D8; line-height:1.6;">
            <p style="margin:0 0 6px; font-weight:600; color:#F0F0F1;">About</p>
            Spec-driven clinical data harmonization.
            SpecRegistry loads domain specs instantly from markdown.
            5-agent pipeline: Ingest → Map → Harmonize → QC → Review.
            LLM-powered fallback resolution, provenance tracking,
            automated QC, and full lineage.
        </div>
        """, unsafe_allow_html=True)

        # API key status
        st.divider()
        st.subheader("API Keys")

        anthropic_configured = bool(os.environ.get("ANTHROPIC_API_KEY"))

        def _key_dot(configured, optional=False):
            if configured:
                return f'<span style="color:{BRAND_TEAL}; font-weight:600;">Connected</span>'
            elif optional:
                return '<span style="color:#6B7280;">Optional</span>'
            else:
                return f'<span style="color:{BRAND_AMBER}; font-weight:600;">Missing</span>'

        st.markdown(
            f'<div style="font-family:Barlow,sans-serif; font-size:0.85rem; line-height:2; color:#C0C8D8;">'
            f'Anthropic: {_key_dot(anthropic_configured, optional=True)}'
            f'</div>',
            unsafe_allow_html=True,
        )

        # Knowledge Base Explorer
        st.divider()
        st.subheader("Knowledge Base")
        st.markdown(
            '<div style="font-family:Barlow,sans-serif; font-size:0.8rem; color:#8B9DC3; '
            'line-height:1.5; margin-bottom:10px;">'
            'The SpecRegistry replaces RAG with deterministic, markdown-based domain '
            'specifications. Each variable has a dedicated spec file defining its '
            'semantic identity, allowed values, mapping rules, and QC checks. '
            'The pipeline loads these instantly at startup&mdash;no vector DB, '
            'no embedding model, no retrieval latency.'
            '</div>',
            unsafe_allow_html=True,
        )

        _render_knowledge_base_tree()

        return {
            "use_llm": use_llm,
            "enable_review": enable_review,
            "domain": domain,
            "skip_qc": skip_qc,
        }


def render_file_upload():
    """Render file upload section."""
    st.subheader("Input Data")

    col1, col2 = st.columns([2, 1])

    with col1:
        uploaded_file = st.file_uploader(
            "Upload source data file",
            type=["csv", "xlsx", "xls", "sas7bdat"],
            help="Supported formats: CSV, Excel, SAS"
        )

    with col2:
        trial_id = st.text_input(
            "Trial ID (optional)",
            placeholder="e.g., NCT12345678",
            help="Leave blank to extract from filename"
        )

    # Optional data dictionary upload
    with st.expander("Data Dictionary (Optional)", expanded=False):
        st.caption("Upload a data dictionary to improve column mapping accuracy")
        data_dict_file = st.file_uploader(
            "Upload data dictionary",
            type=["csv", "xlsx", "xls", "json", "pdf"],
            help="Maps source column names to descriptions or standard terms (tabular PDFs supported)",
            key="data_dict_uploader"
        )

        if data_dict_file:
            st.success(f"Data dictionary loaded: {data_dict_file.name}")

    return uploaded_file, trial_id, data_dict_file if 'data_dict_file' in dir() else None


def render_progress():
    """Render progress indicators with horizontal pipeline stepper."""
    if st.session_state.progress_log:
        latest = st.session_state.progress_log[-1]

        stage_order = ["INGEST", "MAP", "HARMONIZE", "QC", "REVIEW"]
        current_stage = latest["stage"].upper()
        is_error = current_stage == "ERROR"
        is_complete = current_stage in ("COMPLETE", "FINALIZE")

        if is_complete:
            current_idx = len(stage_order)
        elif current_stage in stage_order:
            current_idx = stage_order.index(current_stage)
        else:
            current_idx = 0

        st.markdown(
            pipeline_stepper(stage_order, current_idx, failed=is_error),
            unsafe_allow_html=True
        )

        st.progress(latest["progress"])

        level = "error" if is_error else ("success" if is_complete else "info")
        badge = status_badge(latest["stage"].upper(), level)
        st.markdown(
            f'{badge} <span style="margin-left:10px; font-family:Barlow,sans-serif; '
            f'color:{BRAND_TEXT};">{latest["message"]}</span>',
            unsafe_allow_html=True
        )


def create_results_zip(result: PipelineResult) -> bytes:
    """Create a ZIP file containing all pipeline results.

    Includes:
    - Harmonized data CSV
    - QC report CSV
    - Provenance CSV
    - Transformation report DOCX
    """
    import io
    import zipfile

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        trial_id = result.metadata.get("trial_id", "output")

        # Add harmonized data
        if result.harmonized_data is not None:
            csv_data = result.harmonized_data.to_csv(index=False)
            zf.writestr(f"{trial_id}_harmonized.csv", csv_data)

        # Add QC report
        if result.qc_report is not None and len(result.qc_report) > 0:
            qc_csv = result.qc_report.to_csv(index=False)
            zf.writestr(f"{trial_id}_qc_report.csv", qc_csv)

        # Add provenance data
        if result.provenance_df is not None and len(result.provenance_df) > 0:
            prov_csv = result.provenance_df.to_csv(index=False)
            zf.writestr(f"{trial_id}_provenance.csv", prov_csv)

        # Add transformation report as DOCX
        try:
            docx_report = create_transformation_report_docx(result)
            zf.writestr(f"{trial_id}_transformation_report.docx", docx_report)
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX report")
        except Exception as e:
            logger.warning(f"Could not create DOCX report: {e}")

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


def create_transformation_report_docx(result: PipelineResult) -> bytes:
    """Create a DOCX transformation report. Returns bytes for the DOCX file."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import io

    doc = Document()

    title = doc.add_heading('Harmonization Transformation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Get metadata
    ingest_meta = result.metadata.get('stages', {}).get('ingest', {})
    harmonize_meta = result.metadata.get('stages', {}).get('harmonize', {})

    trial_id = result.metadata.get('trial_id', 'Unknown')
    source_filename = ingest_meta.get('source_filename', result.metadata.get('source_filename', 'Unknown'))
    rows_in = ingest_meta.get('rows', result.metadata.get('rows_processed', 0))
    rows_out = harmonize_meta.get('rows_out', rows_in)
    rows_dropped = rows_in - rows_out if rows_in and rows_out else 0

    doc.add_paragraph(f"Trial: {trial_id}")
    doc.add_paragraph(f"Input: {source_filename}")
    doc.add_paragraph(f"Rows in input: {rows_in} | Rows in output: {rows_out} | Rows dropped: {rows_dropped}")
    doc.add_paragraph(f"Pipeline version: v4 (Spec-Driven Architecture)")
    doc.add_paragraph(f"Run ID: {result.metadata.get('timestamp', 'Unknown')}")

    # Execution Summary
    doc.add_heading('Execution Summary', level=1)
    stages = result.metadata.get('stages', {})
    for stage_name, stage_info in stages.items():
        if isinstance(stage_info, dict):
            status = "✓" if stage_info.get('success', True) else "✗"
            time_ms = stage_info.get('execution_time_ms', 0)
            doc.add_paragraph(f"{status} {stage_name}: {time_ms}ms")

    dict_filename = ingest_meta.get('dictionary_filename') or result.metadata.get('dictionary_filename')
    if dict_filename:
        doc.add_paragraph(f"Dictionary used: {dict_filename}")
    else:
        doc.add_paragraph("Dictionary used: None")

    doc.add_paragraph()

    # Output Schema
    doc.add_heading('1. Output Schema', level=1)
    if result.harmonized_data is not None:
        doc.add_paragraph(', '.join(result.harmonized_data.columns.tolist()))
    else:
        doc.add_paragraph("TRIAL, SUBJID, SEX, RACE, AGE, AGEU, AGEGP, ETHNIC, COUNTRY, SITEID, STUDYID, USUBJID, ARMCD, ARM, BRTHDTC, RFSTDTC, RFENDTC, DOMAIN")

    # Variable Transformation Table
    doc.add_heading('2. Variable-Level Transformations', level=1)

    # Build per-variable confidence summary from provenance for DOCX
    var_confidence_docx = {}
    if result.provenance_df is not None and len(result.provenance_df) > 0 and 'mapping_confidence' in result.provenance_df.columns:
        for var_name, group in result.provenance_df.groupby('variable'):
            counts = group['mapping_confidence'].value_counts().to_dict()
            total = len(group)
            dominant = max(counts, key=counts.get) if counts else "—"
            dist_parts = []
            for level in ["HIGH", "MEDIUM", "LOW", "UNMAPPED"]:
                if level in counts:
                    pct = counts[level] / total * 100
                    dist_parts.append(f"{level}: {pct:.0f}%")
            var_confidence_docx[var_name] = dominant + " (" + ", ".join(dist_parts) + ")" if dist_parts else "—"

    if result.lineage:
        trans_table = doc.add_table(rows=1, cols=8)
        trans_table.style = 'Table Grid'

        headers = ['Variable', 'Source', 'Operation', 'Details', 'Confidence', 'Changed', '%', 'Missing']
        hdr_cells = trans_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        for entry in result.lineage:
            row_cells = trans_table.add_row().cells
            var_name = str(entry.get('variable', ''))
            row_cells[0].text = var_name
            source_col = entry.get('source_column', '') or ''
            row_cells[1].text = str(source_col) if source_col else '(derived)'
            row_cells[2].text = str(entry.get('mapping_operation', 'Copy'))
            transform_op = entry.get('transform_operation', '') or entry.get('transformation', 'None')
            row_cells[3].text = str(transform_op)[:50]
            row_cells[4].text = var_confidence_docx.get(var_name, "—")
            row_cells[5].text = str(entry.get('rows_changed', 0))
            row_cells[6].text = f"{entry.get('percent_changed', 0):.1f}%"
            row_cells[7].text = str(entry.get('missing_count', 0))

    # QC Report Section
    doc.add_heading('3. QC Report', level=1)

    if result.qc_report is not None and len(result.qc_report) > 0:
        qc_table = doc.add_table(rows=1, cols=5)
        qc_table.style = 'Table Grid'

        qc_headers = ['TRIAL', 'Issue Type', 'Variable', 'Rows Affected', 'Notes']
        hdr_cells = qc_table.rows[0].cells
        for i, header in enumerate(qc_headers):
            hdr_cells[i].text = header

        for _, issue in result.qc_report.iterrows():
            row_cells = qc_table.add_row().cells
            row_cells[0].text = str(trial_id)
            row_cells[1].text = str(issue.get('issue_type', ''))
            row_cells[2].text = str(issue.get('variable', ''))
            row_cells[3].text = str(issue.get('n_rows_affected', 0))
            row_cells[4].text = str(issue.get('notes', ''))[:100]
    else:
        doc.add_paragraph("No QC issues found.")

    # Files Produced
    doc.add_heading('4. Files Produced', level=1)
    doc.add_paragraph(f"Harmonized output: {trial_id}_DM_harmonized_*.csv")
    doc.add_paragraph(f"QC report: {trial_id}_QC_report_*.csv")
    doc.add_paragraph(f"Provenance: {trial_id}_provenance_*.csv")
    doc.add_paragraph(f"Transformation report: This document")

    # LLM Review section
    if result.review_result:
        doc.add_heading('5. LLM Review', level=1)

        review = result.review_result
        doc.add_paragraph(f"Approval: {review.get('approval', 'Unknown')}")
        doc.add_paragraph(f"Quality: {review.get('overall_quality', 'Unknown')}")

        if review.get('reason'):
            doc.add_paragraph(f"Summary: {review.get('reason')}")

    # Provenance Summary
    if result.provenance_df is not None and len(result.provenance_df) > 0:
        doc.add_heading('6. Provenance Summary', level=1)
        doc.add_paragraph(f"Total provenance records: {len(result.provenance_df)}")
        if 'mapping_confidence' in result.provenance_df.columns:
            conf_counts = result.provenance_df['mapping_confidence'].value_counts()
            for conf, count in conf_counts.items():
                doc.add_paragraph(f"  {conf}: {count}")

    # Warnings and Errors
    if result.warnings:
        doc.add_heading('Warnings', level=1)
        for w in result.warnings:
            doc.add_paragraph(f"• {w}", style='List Bullet')

    if result.errors:
        doc.add_heading('Errors', level=1)
        for e in result.errors:
            doc.add_paragraph(f"• {e}", style='List Bullet')

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def render_provenance_tab(result: PipelineResult):
    """Render the provenance data tab with summary and detail views."""
    if result.provenance_df is None or len(result.provenance_df) == 0:
        st.info("No provenance data recorded for this run.")
        return

    prov_df = result.provenance_df

    # Summary cards
    total = len(prov_df)
    n_vars = prov_df['variable'].nunique() if 'variable' in prov_df.columns else 0

    conf_counts = {}
    if 'mapping_confidence' in prov_df.columns:
        conf_counts = prov_df['mapping_confidence'].value_counts().to_dict()

    high_ct = conf_counts.get("HIGH", 0)
    medium_ct = conf_counts.get("MEDIUM", 0)
    low_ct = conf_counts.get("LOW", 0)
    unmapped_ct = conf_counts.get("UNMAPPED", 0)

    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.markdown(metric_card("Records", str(total), BRAND_BLUE), unsafe_allow_html=True)
    with col2:
        st.markdown(metric_card("HIGH", str(high_ct), BRAND_TEAL), unsafe_allow_html=True)
    with col3:
        st.markdown(metric_card("MEDIUM", str(medium_ct), BRAND_BLUE), unsafe_allow_html=True)
    with col4:
        st.markdown(metric_card("LOW", str(low_ct), BRAND_AMBER), unsafe_allow_html=True)
    with col5:
        st.markdown(metric_card("UNMAPPED", str(unmapped_ct), BRAND_RED), unsafe_allow_html=True)

    st.markdown("<div style='height:12px;'></div>", unsafe_allow_html=True)

    # Variable filter
    if 'variable' in prov_df.columns:
        variables = ["All"] + sorted(prov_df['variable'].unique().tolist())
        selected_var = st.selectbox("Filter by variable", variables)
        if selected_var != "All":
            prov_df = prov_df[prov_df['variable'] == selected_var]

    # Data table
    st.dataframe(prov_df, use_container_width=True)

    # Download
    csv = prov_df.to_csv(index=False)
    st.download_button(
        "Download Provenance Data",
        csv,
        file_name="provenance.csv",
        mime="text/csv"
    )


def render_results(result: PipelineResult):
    """Render pipeline results with branded card metrics and stoplight panel."""
    st.header("Results")

    # --- Row 1: Summary metric cards ---
    rows = len(result.harmonized_data) if result.harmonized_data is not None else 0
    qc_issues = len(result.qc_report) if result.qc_report is not None else 0
    time_sec = result.execution_time_ms / 1000
    prov_records = len(result.provenance_df) if result.provenance_df is not None else 0

    status_text = "Success" if result.success else "Failed"
    status_color = BRAND_TEAL if result.success else BRAND_RED

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(metric_card("Status", status_text, status_color, value_color=status_color), unsafe_allow_html=True)
    with col2:
        st.markdown(metric_card("Rows Processed", str(rows), BRAND_BLUE), unsafe_allow_html=True)
    with col3:
        qc_color = BRAND_TEAL if qc_issues == 0 else BRAND_AMBER
        st.markdown(metric_card("QC Issues", str(qc_issues), qc_color), unsafe_allow_html=True)
    with col4:
        st.markdown(metric_card("Execution Time", f"{time_sec:.1f}s", BRAND_PURPLE), unsafe_allow_html=True)

    # --- Row 2: LLM + provenance metric cards ---
    llm_enabled = result.metadata.get("llm_enabled", False)
    tokens = result.metadata.get("llm_tokens_used", 0)
    model = result.metadata.get("llm_model", "none")
    review_status = "N/A"
    if result.review_result:
        approval = result.review_result.get("approval", "unknown")
        review_status = approval.replace("_", " ").title()

    col5, col6, col7, col8 = st.columns(4)
    with col5:
        llm_text = "Enabled" if llm_enabled else "Disabled"
        llm_color = BRAND_BLUE if llm_enabled else "#9CA3AF"
        st.markdown(metric_card("LLM", llm_text, llm_color), unsafe_allow_html=True)
    with col6:
        st.markdown(metric_card("Provenance", f"{prov_records:,}", BRAND_PURPLE), unsafe_allow_html=True)
    with col7:
        rev_color = BRAND_TEAL if "green" in review_status.lower() else (BRAND_AMBER if "yellow" in review_status.lower() else BRAND_PURPLE)
        st.markdown(metric_card("Review", review_status, rev_color), unsafe_allow_html=True)
    with col8:
        st.markdown(metric_card("Model", model if model != "none" else "-", BRAND_PURPLE), unsafe_allow_html=True)

    st.markdown("<div style='height:12px;'></div>", unsafe_allow_html=True)

    # --- Errors and warnings ---
    if result.errors:
        st.error("**Errors:**\n" + "\n".join(f"- {e}" for e in result.errors))

    if result.warnings:
        st.warning("**Warnings:**\n" + "\n".join(f"- {w}" for w in result.warnings))

    # --- Download All as ZIP ---
    if result.success:
        zip_data = create_results_zip(result)
        if zip_data:
            trial_id = result.metadata.get("trial_id", "output")
            st.download_button(
                "Download All Results (ZIP)",
                zip_data,
                file_name=f"{trial_id}_harmonization_results.zip",
                mime="application/zip",
                type="primary"
            )

    # --- Tabs for detailed results (added Provenance tab) ---
    tabs = st.tabs(["Harmonized Data", "QC Report", "Provenance", "LLM Review", "Transformations", "Downloads"])

    with tabs[0]:
        if result.harmonized_data is not None:
            st.dataframe(result.harmonized_data, use_container_width=True)
            csv = result.harmonized_data.to_csv(index=False)
            st.download_button(
                "Download Harmonized Data",
                csv,
                file_name="harmonized_data.csv",
                mime="text/csv"
            )
        else:
            st.info("No harmonized data available")

    with tabs[1]:
        if result.qc_report is not None and len(result.qc_report) > 0:
            st.dataframe(result.qc_report, use_container_width=True)
            csv = result.qc_report.to_csv(index=False)
            st.download_button(
                "Download QC Report",
                csv,
                file_name="qc_report.csv",
                mime="text/csv"
            )
        else:
            st.success("No QC issues found.")

    with tabs[2]:
        render_provenance_tab(result)

    with tabs[3]:
        if result.review_result:
            review = result.review_result

            # Branded stoplight panel
            stoplight_level = review.get("stoplight") or review.get("approval", "unknown")
            reason = review.get("reason", "")
            core_present = review.get("core_variables_present", [])
            core_missing = review.get("core_variables_missing", [])

            st.markdown(
                stoplight_panel(stoplight_level, reason, core_present, core_missing or None),
                unsafe_allow_html=True,
            )

            # Quality assessment
            if "overall_quality" in review:
                quality = review["overall_quality"]
                q_lower = quality.lower() if isinstance(quality, str) else ""
                if q_lower in ("good", "excellent"):
                    q_level = "success"
                elif q_lower in ("acceptable", "fair"):
                    q_level = "warning"
                elif q_lower in ("poor", "bad", "unacceptable"):
                    q_level = "error"
                else:
                    q_level = "info"
                st.markdown(
                    f'<p style="margin:12px 0 6px; font-size:0.8rem; color:#6B7280; '
                    f'text-transform:uppercase; letter-spacing:0.05em;">Quality Rating</p>'
                    f'{status_badge(quality.title() if isinstance(quality, str) else str(quality), q_level)}',
                    unsafe_allow_html=True,
                )

            # Formatting issues
            if review.get("formatting_issues"):
                with st.expander("Formatting Issues"):
                    for issue in review["formatting_issues"]:
                        st.markdown(f"- {issue}")

            # Critical issues
            if review.get("critical_issues"):
                with st.expander("Critical Issues"):
                    for issue in review["critical_issues"]:
                        if isinstance(issue, dict):
                            st.markdown(f"- {issue.get('issue', issue)}")
                        else:
                            st.markdown(f"- {issue}")

            # Recommendations
            if review.get("recommendations"):
                with st.expander("Recommendations"):
                    for rec in review["recommendations"]:
                        st.markdown(f"- {rec}")

            # Provenance summary in review context
            if review.get("provenance_summary"):
                with st.expander("Provenance Summary"):
                    st.json(review["provenance_summary"])

            # Full details
            with st.expander("Full Review Details"):
                st.json(review)
        else:
            st.info("No LLM review performed (requires Anthropic API key)")

    with tabs[4]:
        st.subheader("Variable Transformations")
        if result.lineage:
            # Build per-variable confidence summary from provenance
            var_confidence = {}
            if result.provenance_df is not None and len(result.provenance_df) > 0 and 'mapping_confidence' in result.provenance_df.columns:
                for var_name, group in result.provenance_df.groupby('variable'):
                    counts = group['mapping_confidence'].value_counts().to_dict()
                    total = len(group)
                    # Dominant confidence = most frequent
                    dominant = max(counts, key=counts.get) if counts else "—"
                    # Build compact distribution string: "HIGH: 95%, MEDIUM: 5%"
                    dist_parts = []
                    for level in ["HIGH", "MEDIUM", "LOW", "UNMAPPED"]:
                        if level in counts:
                            pct = counts[level] / total * 100
                            dist_parts.append(f"{level}: {pct:.0f}%")
                    var_confidence[var_name] = {
                        "dominant": dominant,
                        "distribution": ", ".join(dist_parts) if dist_parts else "—",
                    }

            lineage_data = []
            for entry in result.lineage:
                var_name = entry.get("variable", "")
                conf_info = var_confidence.get(var_name, {})
                lineage_data.append({
                    "variable": var_name,
                    "source_column": entry.get("source_column", "") or "(derived)",
                    "mapping_operation": entry.get("mapping_operation", ""),
                    "transform_operation": entry.get("transform_operation", "") or entry.get("transformation", ""),
                    "confidence": conf_info.get("dominant", "—"),
                    "confidence_dist": conf_info.get("distribution", "—"),
                    "rows_changed": entry.get("rows_changed", 0),
                    "percent_changed": f"{entry.get('percent_changed', 0):.1f}%",
                    "missing_count": entry.get("missing_count", 0),
                    "non_null_count": entry.get("non_null_count", 0)
                })
            lineage_df = pd.DataFrame(lineage_data)
            st.dataframe(lineage_df, use_container_width=True)
        else:
            st.info("No transformation data available")

    with tabs[5]:
        st.subheader("Downloads")
        col1, col2 = st.columns(2)

        with col1:
            if result.harmonized_data is not None:
                trial_id = result.metadata.get("trial_id", "output")
                csv = result.harmonized_data.to_csv(index=False)
                st.download_button(
                    "Harmonized Data (CSV)",
                    csv,
                    file_name=f"{trial_id}_harmonized.csv",
                    mime="text/csv"
                )

            if result.qc_report is not None and len(result.qc_report) > 0:
                qc_csv = result.qc_report.to_csv(index=False)
                st.download_button(
                    "QC Report (CSV)",
                    qc_csv,
                    file_name=f"{trial_id}_qc_report.csv",
                    mime="text/csv"
                )

            if result.provenance_df is not None and len(result.provenance_df) > 0:
                prov_csv = result.provenance_df.to_csv(index=False)
                st.download_button(
                    "Provenance Data (CSV)",
                    prov_csv,
                    file_name=f"{trial_id}_provenance.csv",
                    mime="text/csv"
                )

        with col2:
            try:
                docx_bytes = create_transformation_report_docx(result)
                st.download_button(
                    "Transformation Report (DOCX)",
                    docx_bytes,
                    file_name=f"{trial_id}_transformation_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.warning(f"Could not generate DOCX report: {e}")

            if result.success:
                zip_data = create_results_zip(result)
                if zip_data:
                    st.download_button(
                        "All Results (ZIP)",
                        zip_data,
                        file_name=f"{trial_id}_harmonization_results.zip",
                        mime="application/zip"
                    )

    # Metadata
    with st.expander("Metadata"):
        st.json(result.metadata)


def run_pipeline(uploaded_file, trial_id: str, config: dict, data_dict_file=None):
    """Run the harmonization pipeline."""
    # Reset progress
    st.session_state.progress_log = []
    st.session_state.pipeline_result = None

    # Read data dictionary if provided
    dict_df = None
    data_dict = None
    if data_dict_file is not None:
        try:
            if data_dict_file.name.endswith('.pdf'):
                data_dict = parse_pdf_dictionary(data_dict_file)
                if data_dict:
                    logger.info(f"Parsed PDF dictionary: {len(data_dict)} variables with codes")
                else:
                    st.info("PDF dictionary loaded but no code mappings were extracted.")
            elif data_dict_file.name.endswith('.csv'):
                dict_df = pd.read_csv(data_dict_file)
            elif data_dict_file.name.endswith('.xlsx'):
                dict_df = pd.read_excel(data_dict_file, engine="openpyxl")
            elif data_dict_file.name.endswith('.xls'):
                dict_df = pd.read_excel(data_dict_file, engine="xlrd")
            elif data_dict_file.name.endswith('.json'):
                data_dict = json.load(data_dict_file)
            logger.info(f"Loaded dictionary file: {data_dict_file.name}")
        except Exception as e:
            st.warning(f"Could not load data dictionary: {e}")
            logger.exception("Data dictionary load error")

    # Read uploaded file
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
        elif uploaded_file.name.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(uploaded_file)
        elif uploaded_file.name.endswith('.sas7bdat'):
            try:
                import pyreadstat
                df, meta = pyreadstat.read_sas7bdat(uploaded_file)
                logger.info(f"Loaded SAS file with {len(df)} rows using pyreadstat")
            except ImportError:
                try:
                    from sas7bdat import SAS7BDAT
                    with SAS7BDAT(uploaded_file) as f:
                        df = f.to_data_frame()
                    logger.info(f"Loaded SAS file with {len(df)} rows using sas7bdat")
                except ImportError:
                    st.error("SAS7BDAT support requires 'pyreadstat' or 'sas7bdat' package. "
                             "Install with: pip install pyreadstat")
                    return None
        else:
            st.error(f"Unsupported file format: {uploaded_file.name}")
            return None
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None

    # Extract trial_id from filename if not provided
    if not trial_id:
        match = re.search(r'(NCT\d{8})', uploaded_file.name)
        trial_id = match.group(1) if match else Path(uploaded_file.name).stem

    # Create progress display elements
    progress_container = st.container()
    with progress_container:
        progress_bar = st.progress(0, text="Initializing pipeline...")
        status_text = st.empty()

    def streamlit_progress_callback(stage: str, status: str, message: str, progress: float):
        """Update Streamlit progress elements in real-time."""
        progress_value = max(0.0, min(1.0, progress))
        progress_bar.progress(progress_value, text=f"{stage.upper()}: {message}")

        st.session_state.progress_log.append({
            "timestamp": datetime.now().isoformat(),
            "stage": stage,
            "status": status,
            "message": message,
            "progress": progress
        })

    # Create orchestrator — v4: spec-driven, domain-parameterized
    reset_settings()
    # Clean domain label (strip preview/beta suffixes from UI display)
    domain_code = config["domain"].split("(")[0].strip().upper()
    orchestrator = create_orchestrator(
        use_llm=config["use_llm"],
        enable_review=config["enable_review"],
        domain=domain_code,
        progress_callback=streamlit_progress_callback,
    )

    # Parse data dictionary using deterministic parser
    if dict_df is not None and data_dict is None:
        progress_bar.progress(0.05, text="INIT: Parsing data dictionary...")

        try:
            data_dict = parse_data_dictionary(dict_df)
            if data_dict:
                st.success(f"Dictionary parsed: {len(data_dict)} variables extracted")
                logger.info(f"Parsed dictionary variables: {list(data_dict.keys())}")
                for var, info in data_dict.items():
                    logger.info(f"  {var}: {list(info.get('codes', {}).keys())}")
            else:
                st.warning("Could not extract code mappings from dictionary")
        except Exception as e:
            logger.exception("Failed to parse dictionary")
            st.warning(f"Dictionary parsing failed: {e}")

    # Run pipeline
    try:
        result = orchestrator.run(
            input_df=df,
            trial_id=trial_id,
            skip_qc=config["skip_qc"],
            data_dict=data_dict
        )
    except Exception as e:
        logger.exception("Pipeline run failed with exception")
        st.error(f"Pipeline error: {str(e)}")
        import traceback
        with st.expander("Error Details"):
            st.code(traceback.format_exc())
        return None

    # Clear progress elements after completion
    if result.success:
        progress_bar.progress(1.0, text="Pipeline complete")
    else:
        progress_bar.progress(1.0, text="Pipeline failed")
        if result.errors:
            for err in result.errors:
                st.error(f"Error: {err}")

    st.session_state.pipeline_result = result
    return result


def main():
    """Main application entry point."""
    init_session_state()

    # Inject brand CSS
    inject_brand_css()

    # Branded header
    st.markdown(f"""
    <div style="display:flex; align-items:center; gap:16px; margin-bottom:4px;">
        <div>
            <h1 style="margin:0; font-size:1.8rem; font-family:Barlow Semi Condensed,sans-serif;
                        font-weight:700; color:{BRAND_NAVY} !important;">Concordia Pipeline v4</h1>
            <p style="margin:2px 0 0; color:{BRAND_BLUE}; font-weight:500;
                      font-family:Barlow,sans-serif; font-size:0.95rem;">
                Spec-Driven Clinical Trial Data Harmonization
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if not PIPELINE_AVAILABLE:
        st.error(f"Pipeline components not available: {IMPORT_ERROR}")
        st.info("Please ensure all dependencies are installed.")
        return

    # Sidebar configuration
    config = render_sidebar()

    # Main content area
    col1, col2 = st.columns([1, 1])

    with col1:
        uploaded_file, trial_id, data_dict_file = render_file_upload()

        if uploaded_file is not None:
            with st.expander("File Preview", expanded=True):
                try:
                    if uploaded_file.name.endswith('.csv'):
                        preview_df = pd.read_csv(uploaded_file, nrows=5)
                    elif uploaded_file.name.endswith(('.xlsx', '.xls')):
                        preview_df = pd.read_excel(uploaded_file, nrows=5)
                    elif uploaded_file.name.endswith('.sas7bdat'):
                        try:
                            import pyreadstat
                            preview_df, _ = pyreadstat.read_sas7bdat(uploaded_file)
                            preview_df = preview_df.head(5)
                        except ImportError:
                            try:
                                from sas7bdat import SAS7BDAT
                                with SAS7BDAT(uploaded_file) as f:
                                    preview_df = f.to_data_frame().head(5)
                            except ImportError:
                                st.warning("Install 'pyreadstat' to preview SAS files")
                                preview_df = None
                    else:
                        preview_df = None

                    if preview_df is not None:
                        st.dataframe(preview_df, use_container_width=True)
                        st.caption(f"Showing first 5 rows | {len(preview_df.columns)} columns")

                    # Reset file position for later reading
                    uploaded_file.seek(0)
                except Exception as e:
                    st.error(f"Error previewing file: {e}")

    with col2:
        st.subheader("Run Pipeline")

        if uploaded_file is not None:
            if st.button("Start Harmonization", type="primary", use_container_width=True):
                result = run_pipeline(uploaded_file, trial_id, config, data_dict_file)
        else:
            st.info("Upload a file to begin")

        # Stage indicators — always visible in the run section
        stage_order = ["INGEST", "MAP", "HARMONIZE", "QC", "REVIEW"]

        if st.session_state.pipeline_result is not None:
            # Pipeline finished — show final state
            st.markdown(
                pipeline_stepper(
                    stage_order,
                    len(stage_order),
                    failed=not st.session_state.pipeline_result.success,
                ),
                unsafe_allow_html=True,
            )
            status_label = "Pipeline complete" if st.session_state.pipeline_result.success else "Pipeline failed"
            st.progress(1.0, text=status_label)
        elif st.session_state.progress_log:
            # Pipeline running — show live progress
            render_progress()
        else:
            # No run yet — show pending stepper
            st.markdown(
                pipeline_stepper(stage_order, -1),
                unsafe_allow_html=True,
            )

    # Results section
    st.divider()

    if st.session_state.pipeline_result is not None:
        render_results(st.session_state.pipeline_result)


if __name__ == "__main__":
    main()
